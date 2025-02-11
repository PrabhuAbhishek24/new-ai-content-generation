import openai
import os
import zipfile
import io
from fpdf import FPDF
from docx import Document
from pathlib import Path
from dotenv import load_dotenv
from docx.shared import Inches
from fastapi import FastAPI, HTTPException
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from typing import Literal


session_data = {}

app = FastAPI()


class UserTypeInput(BaseModel):
    user_type: str  # "free" or "paid"

class DomainInput(BaseModel):
    domain: str  # The domain selected by the user

# Request Models
class GenerateRequest(BaseModel):
    query: str

class QueryInput(BaseModel):
    query: str  # The actual query from the user


class DownloadRequest(BaseModel):
    response: str
    format: Literal["pdf_scorm", "docx_scorm"]

openai.api_key=os.getenv("OPENAI_API_KEY")

def fetch_gpt_response(domain: str, query: str, token_limit: int):
    try:
        system_prompt = (
            f"You are an expert in the {domain} domain only. "
            f"Only answer the questions related to the specified {domain} domain "
            "and don't answer any other questions."
        )
        
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": query},
            ],
            max_tokens=token_limit  # Set token limit based on user type
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"

    

def save_as_scorm_pdf(content, output_folder="scorm_package", scorm_zip_name="scorm_package.zip"):
    # Step 1: Create the SCORM folder structure
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Save the PDF
    pdf_file_path = os.path.join(output_folder, "content.pdf")
    save_as_pdf(content, pdf_file_path)


    # Step 2: Create the HTML file
    html_file_path = os.path.join(output_folder, "index.html")
    with open(html_file_path, "w", encoding="utf-8") as html_file:
        html_file.write(f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>SCORM Content</title>
        </head>
        <body>
            <h1>Research Content Response</h1>
            <iframe src="content.pdf" width="100%" height="600px"></iframe>
        </body>
        </html>
        """)

    # Step 3: Create the imsmanifest.xml file
    manifest_file_path = os.path.join(output_folder, "imsmanifest.xml")
    with open(manifest_file_path, "w", encoding="utf-8") as manifest_file:
        manifest_file.write(f"""
        <?xml version="1.0" encoding="UTF-8"?>
        <manifest xmlns="http://www.imsglobal.org/xsd/imscp_v1p1"
                  xmlns:adlcp="http://www.adlnet.org/xsd/adlcp_v1p3"
                  xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
                  xsi:schemaLocation="http://www.imsglobal.org/xsd/imscp_v1p1
                                      http://www.imsglobal.org/xsd/imscp_v1p1.xsd
                                      http://www.adlnet.org/xsd/adlcp_v1p3
                                      http://www.adlnet.org/xsd/adlcp_v1p3.xsd">
            <metadata>
                <schema>ADL SCORM</schema>
                <schemaversion>1.2</schemaversion>
            </metadata>
            <organizations>
                <organization identifier="ORG-1">
                    <title>Research Content</title>
                    <item identifier="ITEM-1" identifierref="RES-1">
                        <title>Research Content Response</title>
                    </item>
                </organization>
            </organizations>
            <resources>
                <resource identifier="RES-1" type="webcontent" href="index.html">
                    <file href="index.html"/>
                    <file href="content.pdf"/>
                </resource>
            </resources>
        </manifest>
        """)

    # Step 4: Zip the SCORM package
    with zipfile.ZipFile(scorm_zip_name, 'w', zipfile.ZIP_DEFLATED) as scorm_zip:
        for foldername, subfolders, filenames in os.walk(output_folder):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, output_folder)
                scorm_zip.write(file_path, arcname)

    # Provide the download button for the SCORM package
    with open(scorm_zip_name, 'rb') as scorm_file:
        return scorm_file.read()


def save_as_pdf(content, file_name="response.pdf"):
    pdf = FPDF()
    pdf.add_page()

    # Add the logo
    pdf.image("assets/logo.jpeg", x=10, y=8, w=30)

    # Title of the document
    pdf.set_font("Arial", style='B', size=16)
    pdf.ln(30)
    pdf.cell(200, 10, txt="Research Content Response", ln=True, align='C')
    pdf.ln(10)

    # Add content
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(190, 10, content)

    # Save the PDF
    pdf.output(file_name)


def save_as_scorm_word(content, file_name="scorm_package.zip"):
    # Create an in-memory zip file
    scorm_zip = io.BytesIO()

    with zipfile.ZipFile(scorm_zip, 'w') as zf:
        # Create and add manifest.xml
        manifest_content = """<manifest>
            <metadata>
                <schema>ADL SCORM</schema>
                <schemaversion>1.2</schemaversion>
            </metadata>
            <resources>
                <resource identifier="res1" type="webcontent" href="response.docx">
                    <file href="response.docx"/>
                    <file href="response.html"/>
                </resource>
            </resources>
        </manifest>"""
        zf.writestr("imanifest.xml", manifest_content)

        # Create DOCX file
        docx_buffer = io.BytesIO()
        doc = Document()
        # Add the logo to the Word document
        logo_path = "assets/logo.jpeg"
        if Path(logo_path).is_file():
            doc.add_picture(logo_path, width=Inches(1.5))
        doc.add_paragraph('\n')
        doc.add_paragraph("Research Content Response", style='Heading 1')
        doc.add_paragraph('\n')
        doc.add_paragraph(content)
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        zf.writestr("response.docx", docx_buffer.getvalue())
           

        html_content = """
       <html>
       <head><title>Research Content Response</title></head>
       <body>
       <h1>Research Content Response</h1>
       <p>{}</p>
       </body>
       </html>
       """.format(content.replace("\n", "<br>"))
       zf.writestr("index.html", html_content)
   

    scorm_zip.seek(0)
    return scorm_zip.getvalue()

@app.post("/select_user")
def select_user(user_input: UserTypeInput):
    """ API to select user type (Free or Paid) """
    user_type = user_input.user_type.lower()
    if user_type not in ["free", "paid"]:
        raise HTTPException(status_code=400, detail="Invalid user type. Choose 'free' or 'paid'.")

    session_data["user_type"] = user_type
    return {"message": f"User type set to {user_type}"}

@app.post("/set_domain")
def set_domain(domain_input: DomainInput):
    """ API to set the query domain """
    domain = domain_input.domain.strip().lower()
    if not domain:
        raise HTTPException(status_code=400, detail="Domain cannot be empty.")

    session_data["domain"] = domain
    return {"message": f"Domain set to {domain}"}

@app.post("/query")
def query_gpt(query_input: QueryInput):
    """ Unified API to query GPT (Handles both free and paid users) """
    if "user_type" not in session_data:
        raise HTTPException(status_code=403, detail="Please select user type first.")

    if "domain" not in session_data:
        raise HTTPException(status_code=400, detail="Please set a domain before querying.")

    user_type = session_data["user_type"]
    domain = session_data["domain"]
    token_limit = 250 if user_type == "free" else 4000

    # Fetch response from GPT
    response = fetch_gpt_response(domain, query_input.query, token_limit)

    # Store generated content in session data
    session_data["generated_content"] = response  

    return {"response": response, "token_limit": token_limit}

@app.post("/api/download-content")
async def download_content(request: DownloadRequest):
    """ API to handle content download (SCORM format) """

    if "user_type" not in session_data:
        raise HTTPException(status_code=403, detail="Please select user type first.")

    if "domain" not in session_data:
        raise HTTPException(status_code=400, detail="Please set a domain before downloading content.")

    if "generated_content" not in session_data:
        raise HTTPException(status_code=400, detail="No generated content found. Please query first.")

    try:
        generated_content = session_data["generated_content"]  # Retrieve stored content

        if request.format == "pdf_scorm":
            scorm_file = save_as_scorm_pdf(generated_content, "pdf")
            return StreamingResponse(io.BytesIO(scorm_file), media_type="application/zip", headers={
                "Content-Disposition": "attachment; filename=scorm_package_pdf.zip"
            })
        elif request.format == "docx_scorm":
            scorm_file = save_as_scorm_word(generated_content, "docx")
            return StreamingResponse(io.BytesIO(scorm_file), media_type="application/zip", headers={
                "Content-Disposition": "attachment; filename=scorm_package_doc.zip"
            })
        else:
            raise HTTPException(status_code=400, detail="Invalid download format selected.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error in downloading content: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
